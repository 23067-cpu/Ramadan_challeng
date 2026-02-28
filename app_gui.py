"""
SupNum - S3C'1447 Challenge Ramadan
Fichier : app_gui.py
Description : Fichier développé en Python pour la résolution du problème RCPSP.
Tous les algorithmes et logiques internes sont optimisés pour des performances maximales.
"""
"""
app_gui.py
Main GUI for RCPSP Solver - SupNum Coding Challenge S3C'1447
Modern dark-themed Tkinter application with:
  - Tab 1: Solver configuration + dataset browser
  - Tab 2: Gantt chart visualization
  - Tab 3: Results comparison table
"""

from __future__ import annotations
import os
import sys
import threading
import time
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from typing import Dict, List, Optional, Tuple
import math

try:
    import matplotlib
    matplotlib.use('TkAgg')
    import matplotlib.pyplot as plt
    import matplotlib.patches as mpatches
    from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg, NavigationToolbar2Tk
    from matplotlib.figure import Figure
    HAS_MATPLOTLIB = True
except ImportErreur:
    HAS_MATPLOTLIB = False

from rcpsp_parser import RCPSPInstance, parse_instance, load_dataset, parse_solutions_file
from rcpsp_solver import HybridGASolver, GAConfig
from rcpsp_bounds import compute_lower_bound
from rcpsp_scheduler import Schedule
from fast_j60_solver import solve as fast_solve


# ─────────────────────────────────────────────────────────────────────────────
# Colour Palette
# ─────────────────────────────────────────────────────────────────────────────
BG_DARK   = "#0f1117"
BG_CARD   = "#1a1d27"
BG_PANEL  = "#1e2130"
ACCENT    = "#6c63ff"
ACCENT2   = "#00d4aa"
ACCENT3   = "#ff6b6b"
TEXT_MAIN = "#e8eaf6"
TEXT_DIM  = "#7b82a8"
BORDER    = "#2d3055"
SUCCESS   = "#43a047"
WARNING   = "#ffa726"
DANGER    = "#ef5350"

FONT_TITLE  = ("Segoe UI", 18, "bold")
FONT_H2     = ("Segoe UI", 13, "bold")
FONT_BODY   = ("Segoe UI", 10)
FONT_MONO   = ("Consolas", 9)
FONT_SMALL  = ("Segoe UI", 8)

# Couleurs des barres de Gantt per job (cycling)
GANTT_COLORS = [
    "#6c63ff", "#00d4aa", "#ff6b6b", "#ffa726", "#42a5f5",
    "#ab47bc", "#26c6da", "#d4e157", "#ff7043", "#ec407a",
    "#26a69a", "#78909c", "#66bb6a", "#ffee58", "#8d6e63",
]


# ─────────────────────────────────────────────────────────────────────────────
# Result record
# ─────────────────────────────────────────────────────────────────────────────
from dataclasses import dataclass, field

@dataclass
class SolverResult:
    instance_name: str
    n_jobs: int
    our_cmax: int
    lb: int
    ref_lb: Optional[int]
    ref_ub: Optional[int]
    gap_to_lb: float            # (our_cmax - lb) / lb * 100
    gap_to_ub: float            # (our_cmax - ref_ub) / ref_ub * 100  (if known)
    elapsed: float
    schedule: Optional[Schedule]
    beats_known: bool = False   # True if our_cmax < ref_ub


# ─────────────────────────────────────────────────────────────────────────────
# Application
# ─────────────────────────────────────────────────────────────────────────────
class RCPSPApp(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("RCPSP Solver — SupNum S3C'1447")
        self.geometry("1280x820")
        self.minsize(960, 640)
        self.configure(bg=BG_DARK)

        # État de l'application
        self.dataset_folder: str = ""
        self.instance_files: List[str] = []
        self.instances: Dict[str, RCPSPInstance] = {}
        self.solutions_ref: Dict[str, Tuple[int, int]] = {}   # name -> (lb, ub)
        self.results: List[SolverResult] = []
        self.current_schedule: Optional[Schedule] = None
        self._solver_thread: Optional[threading.Thread] = None
        self._stop_event = threading.Event()
        self._selected_instance: Optional[str] = None

        self._setup_styles()
        self._build_ui()
        self.protocol("WM_DELETE_WINDOW", self._on_close)

    # ─── Styles ───────────────────────────────────────────────────────────

    def _setup_styles(self):
        style = ttk.Style(self)
        style.theme_use("clam")

        style.configure(".",
                        background=BG_DARK, foreground=TEXT_MAIN,
                        font=FONT_BODY, borderwidth=0)
        style.configure("TNotebook",
                        background=BG_DARK, borderwidth=0, tabmargins=[0, 0, 0, 0])
        style.configure("TNotebook.Tab",
                        background=BG_CARD, foreground=TEXT_DIM,
                        padding=[18, 8], font=("Segoe UI", 10, "bold"),
                        borderwidth=0)
        style.map("TNotebook.Tab",
                  background=[("selected", BG_PANEL)],
                  foreground=[("selected", ACCENT)])
        style.configure("TFrame", background=BG_DARK)
        style.configure("Card.TFrame", background=BG_CARD)
        style.configure("TLabel", background=BG_DARK, foreground=TEXT_MAIN)
        style.configure("Dim.TLabel", background=BG_CARD, foreground=TEXT_DIM)
        style.configure("TButton",
                        background=ACCENT, foreground="white",
                        padding=[12, 6], relief="flat", font=("Segoe UI", 10, "bold"))
        style.map("TButton",
                  background=[("active", "#8078ff"), ("pressed", "#5048cc")])
        style.configure("Danger.TButton",
                        background=DANGER, foreground="white",
                        padding=[12, 6], relief="flat", font=("Segoe UI", 10, "bold"))
        style.configure("TEntry",
                        fieldbackground=BG_PANEL, background=BG_PANEL,
                        foreground=TEXT_MAIN, insertcolor=TEXT_MAIN,
                        relief="flat", padding=[6, 4])
        style.configure("TSpinbox",
                        fieldbackground=BG_PANEL, background=BG_PANEL,
                        foreground=TEXT_MAIN, arrowcolor=TEXT_DIM)
        style.configure("TProgressionbar",
                        troughcolor=BG_PANEL, background=ACCENT,
                        thickness=6)
        style.configure("Treeview",
                        background=BG_CARD, foreground=TEXT_MAIN,
                        fieldbackground=BG_CARD, rowheight=26,
                        font=FONT_MONO)
        style.configure("Treeview.Heading",
                        background=BG_PANEL, foreground=ACCENT,
                        font=("Segoe UI", 9, "bold"), relief="flat")
        style.map("Treeview", background=[("selected", ACCENT)])
        style.configure("TScrollbar",
                        background=BG_PANEL, troughcolor=BG_DARK,
                        arrowcolor=TEXT_DIM, borderwidth=0)
        style.configure("TCheckbutton",
                        background=BG_CARD, foreground=TEXT_MAIN)
        style.configure("TScale",
                        background=BG_CARD, troughcolor=BG_PANEL,
                        sliderthickness=12)

    # ─── UI Build ─────────────────────────────────────────────────────────

    def _build_ui(self):
        # ── Header ──
        header = tk.Frame(self, bg=BG_DARK, height=64)
        header.pack(fill="x", side="top")
        header.pack_propagate(False)

        tk.Label(header, text="⚡ RCPSP Solver", font=("Segoe UI", 16, "bold"),
                 bg=BG_DARK, fg=ACCENT).pack(side="left", padx=20, pady=14)
        tk.Label(header, text="SupNum S3C'1447 — Hybrid GA + Critical-Path LS + Tabu Search + LNS",
                 font=("Segoe UI", 9), bg=BG_DARK, fg=TEXT_DIM).pack(side="left", pady=18)

        tk.Label(header, text="🎓 RCPSP | PSPLIB Benchmark", font=FONT_SMALL,
                 bg=BG_DARK, fg=TEXT_DIM).pack(side="right", padx=20)

        # Separator
        tk.Frame(self, bg=BORDER, height=1).pack(fill="x")

        # ── Notebook ──
        self.notebook = ttk.Notebook(self)
        self.notebook.pack(fill="both", expand=True, padx=6, pady=6)

        self.tab_solver = ttk.Frame(self.notebook)
        self.tab_gantt  = ttk.Frame(self.notebook)
        self.tab_results = ttk.Frame(self.notebook)

        self.notebook.add(self.tab_solver,  text="  🔧 Solver  ")
        self.notebook.add(self.tab_gantt,   text="  📊 Gantt Chart  ")
        self.notebook.add(self.tab_results, text="  📋 Results  ")

        self._build_solver_tab()
        self._build_gantt_tab()
        self._build_results_tab()

    # ─── Tab 1: Solver ────────────────────────────────────────────────────

    def _build_solver_tab(self):
        tab = self.tab_solver
        tab.configure(style="TFrame")

        # Two-column layout
        left = tk.Frame(tab, bg=BG_DARK, width=380)
        left.pack(side="left", fill="y", padx=(8, 4), pady=8)
        left.pack_propagate(False)

        right = tk.Frame(tab, bg=BG_DARK)
        right.pack(side="left", fill="both", expand=True, padx=(4, 8), pady=8)

        # ── Left: Dataset & Config ──
        self._build_dataset_panel(left)
        self._build_config_panel(left)

        # ── Right: Instance list + Log ──
        self._build_instance_panel(right)
        self._build_log_panel(right)

    def _build_dataset_panel(self, parent):
        card = self._make_card(parent, "📁 Dataset")
        card.pack(fill="x", pady=(0, 8))

        # Folder path
        path_frame = tk.Frame(card, bg=BG_CARD)
        path_frame.pack(fill="x", pady=(0, 6))
        self.folder_var = tk.StringVar(value="No folder selected")
        tk.Label(path_frame, textvariable=self.folder_var, bg=BG_CARD,
                 fg=TEXT_DIM, font=FONT_SMALL, anchor="w",
                 wraplength=280).pack(fill="x")

        btn_row = tk.Frame(card, bg=BG_CARD)
        btn_row.pack(fill="x")
        self._btn(btn_row, "Browse Instances Folder", self._browse_folder).pack(side="left")
        self._btn(btn_row, "Load Solutions File", self._browse_solutions, color=ACCENT2).pack(side="left", padx=(6, 0))

        self.sol_status_var = tk.StringVar(value="No solutions file loaded")
        tk.Label(card, textvariable=self.sol_status_var, bg=BG_CARD,
                 fg=TEXT_DIM, font=FONT_SMALL).pack(anchor="w", pady=(4, 0))

        # Instance count
        self.inst_count_var = tk.StringVar(value="0 instances loaded")
        tk.Label(card, textvariable=self.inst_count_var, bg=BG_CARD,
                 fg=ACCENT2, font=("Segoe UI", 9, "bold")).pack(anchor="w")

    def _build_config_panel(self, parent):
        card = self._make_card(parent, "⚙️ Algorithm Parameters")
        card.pack(fill="x", pady=(0, 8))

        params = [
            ("Population Size",    "pop_size",   100,  10, 500, 10),
            ("Generations",        "max_gen",    500,  10,5000, 50),
            ("Time Limit (sec)",   "time_limit", 120,   5, 600,  5),
            ("Crossover Rate",     "crossover",  0.85, 0.1, 1.0, 0.05),
            ("Mutation Rate",      "mutation",   0.05, 0.01, 0.5, 0.01),
            ("Tournament Size",    "tournament", 4,    2,   10,   1),
            ("LS Iterations",      "ls_iters",   150,  10, 500,  10),
            ("Tabu Tenure",        "tabu_tenure",12,    5,  50,   1),
            ("Stagnation Limit",   "stagnation", 80,   10, 200,  10),
        ]

        self._param_vars: Dict[str, tk.Variable] = {}
        for label, key, default, mn, mx, step in params:
            row = tk.Frame(card, bg=BG_CARD)
            row.pack(fill="x", pady=2)
            tk.Label(row, text=label, bg=BG_CARD, fg=TEXT_MAIN,
                     font=FONT_BODY, width=20, anchor="w").pack(side="left")

            if isinstance(default, float):
                var = tk.DoubleVar(value=default)
                sp = tk.Spinbox(row, from_=mn, to=mx, increment=step,
                                textvariable=var, width=8,
                                bg=BG_PANEL, fg=TEXT_MAIN, relief="flat",
                                buttonbackground=BG_PANEL, insertbackground=TEXT_MAIN)
            else:
                var = tk.IntVar(value=default)
                sp = tk.Spinbox(row, from_=mn, to=mx, increment=step,
                                textvariable=var, width=8,
                                bg=BG_PANEL, fg=TEXT_MAIN, relief="flat",
                                buttonbackground=BG_PANEL, insertbackground=TEXT_MAIN)
            sp.pack(side="right")
            self._param_vars[key] = var

        # Options
        sep = tk.Frame(card, bg=BORDER, height=1)
        sep.pack(fill="x", pady=6)

        self.use_psgs_var = tk.BooleanVar(value=True)
        self.use_res_prio_var = tk.BooleanVar(value=True)
        self.use_dual_decode_var = tk.BooleanVar(value=True)
        self.use_lns_var = tk.BooleanVar(value=True)

        opts = tk.Frame(card, bg=BG_CARD)
        opts.pack(fill="x")
        for text, var in [
            ("Dual Decode (SSGS + PSGS, keep best)", self.use_dual_decode_var),
            ("Large Neighbourhood Search (LNS)",    self.use_lns_var),
            ("Adaptive Mutation + Diversity Check",  self.use_res_prio_var),
        ]:
            tk.Checkbutton(opts, text=text, variable=var,
                           bg=BG_CARD, fg=TEXT_MAIN, selectcolor=BG_PANEL,
                           activebackground=BG_CARD, activeforeground=ACCENT,
                           font=FONT_SMALL).pack(anchor="w")

    def _build_instance_panel(self, parent):
        card = self._make_card(parent, "📌 Instances")
        card.pack(fill="both", expand=True, pady=(0, 4))

        # Search bar
        search_row = tk.Frame(card, bg=BG_CARD)
        search_row.pack(fill="x", pady=(0, 6))
        tk.Label(search_row, text="🔍", bg=BG_CARD, fg=TEXT_DIM).pack(side="left", padx=(0, 4))
        self.search_var = tk.StringVar()
        self.search_var.trace_add("write", self._filter_instances)
        search_entry = tk.Entry(search_row, textvariable=self.search_var,
                                bg=BG_PANEL, fg=TEXT_MAIN, relief="flat",
                                insertbackground=TEXT_MAIN, font=FONT_BODY)
        search_entry.pack(fill="x", expand=True)

        # Listbox with scrollbar
        lf = tk.Frame(card, bg=BG_CARD)
        lf.pack(fill="both", expand=True)
        sb = ttk.Scrollbar(lf, orient="vertical")
        self.instance_listbox = tk.Listbox(
            lf, bg=BG_PANEL, fg=TEXT_MAIN, selectbackground=ACCENT,
            relief="flat", font=FONT_MONO, yscrollcommand=sb.set,
            activestyle="none", highlightthickness=0, borderwidth=0)
        sb.config(command=self.instance_listbox.yview)
        sb.pack(side="right", fill="y")
        self.instance_listbox.pack(fill="both", expand=True)
        self.instance_listbox.bind("<<ListboxSelect>>", self._on_instance_select)

        # Action buttons
        btn_row = tk.Frame(card, bg=BG_CARD)
        btn_row.pack(fill="x", pady=(6, 0))
        self._btn(btn_row, "▶ Run Selected", self._run_selected).pack(side="left")
        self._btn(btn_row, "▶▶ Run All Open", self._run_all_open,
                  color=ACCENT2).pack(side="left", padx=6)
        self._btn(btn_row, "⏹ Stop", self._stop_solver,
                  color=DANGER).pack(side="left")

        # Progression
        self.progress_var = tk.DoubleVar(value=0)
        self.progress_label_var = tk.StringVar(value="Ready")
        tk.Label(card, textvariable=self.progress_label_var,
                 bg=BG_CARD, fg=ACCENT2, font=FONT_SMALL).pack(anchor="w", pady=(4, 0))
        ttk.Progressionbar(card, variable=self.progress_var, maximum=100,
                        style="TProgressionbar").pack(fill="x", pady=(2, 0))

    def _build_log_panel(self, parent):
        card = self._make_card(parent, "📝 Log")
        card.pack(fill="x", pady=(4, 0))

        log_frame = tk.Frame(card, bg=BG_CARD)
        log_frame.pack(fill="both", expand=True)
        vsb = ttk.Scrollbar(log_frame, orient="vertical")
        self.log_text = tk.Text(
            log_frame, bg=BG_PANEL, fg=TEXT_MAIN, relief="flat",
            font=FONT_MONO, height=8, wrap="word",
            yscrollcommand=vsb.set, state="disabled",
            insertbackground=TEXT_MAIN, highlightthickness=0)
        vsb.config(command=self.log_text.yview)
        vsb.pack(side="right", fill="y")
        self.log_text.pack(fill="both", expand=True)

        self.log_text.tag_config("info",    foreground=TEXT_MAIN)
        self.log_text.tag_config("success", foreground=ACCENT2)
        self.log_text.tag_config("warning", foreground=WARNING)
        self.log_text.tag_config("error",   foreground=DANGER)
        self.log_text.tag_config("accent",  foreground=ACCENT)

        btn_row = tk.Frame(card, bg=BG_CARD)
        btn_row.pack(fill="x", pady=(4, 0))
        self._btn(btn_row, "Clear Log", self._clear_log, color=TEXT_DIM).pack(side="right")

    # ─── Tab 2: Gantt Chart ───────────────────────────────────────────────

    def _build_gantt_tab(self):
        tab = self.tab_gantt
        tab.configure(style="TFrame")

        if not HAS_MATPLOTLIB:
            tk.Label(tab,
                     text="⚠️ matplotlib not installed.\nRun: pip install matplotlib",
                     bg=BG_DARK, fg=DANGER, font=FONT_H2).pack(expand=True)
            return

        # Controls bar
        ctrl = tk.Frame(tab, bg=BG_DARK)
        ctrl.pack(fill="x", padx=8, pady=6)
        tk.Label(ctrl, text="Gantt Chart — Best Schedule",
                 font=FONT_H2, bg=BG_DARK, fg=TEXT_MAIN).pack(side="left")
        self._btn(ctrl, "🔄 Refresh", self._refresh_gantt,
                  color=ACCENT2).pack(side="right")
        self._btn(ctrl, "💾 Save PNG", self._save_gantt_png,
                  color=TEXT_DIM).pack(side="right", padx=6)

        # Gantt info labels
        info_frame = tk.Frame(tab, bg=BG_DARK)
        info_frame.pack(fill="x", padx=8)
        self.gantt_info_var = tk.StringVar(value="No schedule to display. Run the solver first.")
        tk.Label(info_frame, textvariable=self.gantt_info_var,
                 bg=BG_DARK, fg=TEXT_DIM, font=FONT_BODY).pack(side="left")

        # Matplotlib figure
        self.gantt_fig = Figure(figsize=(12, 6), dpi=90, facecolor=BG_CARD)
        self.gantt_ax = self.gantt_fig.add_subplot(111)
        self._style_axes(self.gantt_ax, "No schedule loaded")

        self.gantt_canvas = FigureCanvasTkAgg(self.gantt_fig, master=tab)
        self.gantt_canvas.get_tk_widget().pack(fill="both", expand=True, padx=8, pady=4)

        toolbar_frame = tk.Frame(tab, bg=BG_DARK)
        toolbar_frame.pack(fill="x", padx=8)
        toolbar = NavigationToolbar2Tk(self.gantt_canvas, toolbar_frame)
        toolbar.config(background=BG_DARK)
        toolbar.update()

    # ─── Tab 3: Results ───────────────────────────────────────────────────

    def _build_results_tab(self):
        tab = self.tab_results
        tab.configure(style="TFrame")

        # Header row
        hdr = tk.Frame(tab, bg=BG_DARK)
        hdr.pack(fill="x", padx=8, pady=6)
        tk.Label(hdr, text="Results vs PSPLIB Benchmark",
                 font=FONT_H2, bg=BG_DARK, fg=TEXT_MAIN).pack(side="left")
        self._btn(hdr, "📥 Export Excel", self._export_excel,
                  color=ACCENT2).pack(side="right")
        self._btn(hdr, "📄 Export Word", self._export_word,
                  color=ACCENT).pack(side="right", padx=6)
        self._btn(hdr, "🗑 Clear Results", self._clear_results,
                  color=DANGER).pack(side="right", padx=6)

        # Summary bar
        self.summary_frame = tk.Frame(tab, bg=BG_DARK)
        self.summary_frame.pack(fill="x", padx=8, pady=(0, 6))
        self._update_summary_bar()

        # Treeview
        cols = ("Instance", "Jobs", "Our Cmax", "LB", "Ref LB", "Ref UB",
                "Gap to LB%", "Gap to UB%", "Time(s)", "Status")
        tree_frame = tk.Frame(tab, bg=BG_DARK)
        tree_frame.pack(fill="both", expand=True, padx=8, pady=(0, 8))

        vsb = ttk.Scrollbar(tree_frame, orient="vertical")
        hsb = ttk.Scrollbar(tree_frame, orient="horizontal")

        self.results_tree = ttk.Treeview(
            tree_frame, columns=cols, show="headings",
            yscrollcommand=vsb.set, xscrollcommand=hsb.set,
            style="Treeview")
        vsb.config(command=self.results_tree.yview)
        hsb.config(command=self.results_tree.xview)

        col_widths = [120, 55, 90, 65, 75, 75, 100, 100, 75, 100]
        for col, w in zip(cols, col_widths):
            self.results_tree.heading(col, text=col)
            self.results_tree.column(col, width=w, anchor="center")

        self.results_tree.tag_configure("beat",    background="#1a3320", foreground=ACCENT2)
        self.results_tree.tag_configure("good",    background="#1a2b1a", foreground=SUCCESS)
        self.results_tree.tag_configure("neutral", background=BG_CARD,   foreground=TEXT_MAIN)
        self.results_tree.tag_configure("open",    background="#2a1a2a", foreground=WARNING)

        vsb.pack(side="right", fill="y")
        hsb.pack(side="bottom", fill="x")
        self.results_tree.pack(fill="both", expand=True)

    # ─── Actions ──────────────────────────────────────────────────────────

    def _browse_folder(self):
        folder = filedialog.askdirectory(title="Select PSPLIB Instance Folder")
        if not folder:
            return
        self.dataset_folder = folder
        self.folder_var.set(os.path.basename(folder) + "/")
        self._load_instances(folder)

    def _browse_solutions(self):
        path = filedialog.askopenfilename(
            title="Select Solutions File (e.g. j30hrs.sm)",
            filetypes=[("Solution files", "*.sm *.txt *.csv"), ("All files", "*.*")])
        if not path:
            return
        self.solutions_ref = parse_solutions_file(path)
        n = len(self.solutions_ref)
        self.sol_status_var.set(f"✓ Loaded {n} reference solutions from {os.path.basename(path)}")
        self._log(f"Loaded {n} reference solutions from {os.path.basename(path)}", "success")

    def _load_instances(self, folder: str):
        self.instance_files = []
        self.instances = {}
        for fname in sorted(os.listdir(folder)):
            if fname.endswith('.sm') or fname.endswith('.SM'):
                fpath = os.path.join(folder, fname)
                self.instance_files.append(fpath)
        total = len(self.instance_files)
        self.inst_count_var.set(f"{total} instance files found")
        self._log(f"Found {total} .sm files in {os.path.basename(folder)}", "accent")
        self._refresh_instance_list()

    def _refresh_instance_list(self):
        self.instance_listbox.delete(0, tk.END)
        q = self.search_var.get().lower()
        for fpath in self.instance_files:
            name = os.path.basename(fpath)
            if q and q not in name.lower():
                continue
            self.instance_listbox.insert(tk.END, f"  {name}")

    def _filter_instances(self, *_):
        self._refresh_instance_list()

    def _on_instance_select(self, event):
        sel = self.instance_listbox.curselection()
        if not sel:
            return
        text = self.instance_listbox.get(sel[0]).strip()
        for fpath in self.instance_files:
            if os.path.basename(fpath) == text:
                self._selected_instance = fpath
                self._log(f"Selected: {text}", "info")
                break

    def _run_selected(self):
        if not self._selected_instance:
            messagebox.showwarning("No selection", "Please select an instance from the list.")
            return
        self._run_instances([self._selected_instance])

    def _run_all_open(self):
        if not self.instance_files:
            messagebox.showwarning("No instances", "Please load a dataset folder first.")
            return
        self._run_instances(list(self.instance_files))

    def _run_instances(self, files: List[str]):
        if self._solver_thread and self._solver_thread.is_alive():
            messagebox.showwarning("Running", "Solver is already running. Stop it first.")
            return
        self._stop_event.clear()
        self._solver_thread = threading.Thread(
            target=self._solver_worker, args=(files,), daemon=True)
        self._solver_thread.start()

    def _stop_solver(self):
        self._stop_event.set()
        self._log("⏹ Stop signal sent. Finishing current instance...", "warning")

    def _solver_worker(self, files: List[str]):
        total = len(files)
        self._log(f"🚀 Démarrage du solveur on {total} instance(s)...", "accent")
        self.progress_var.set(0)

        config = GAConfig(
            pop_size=int(self._param_vars["pop_size"].get()),
            max_generations=int(self._param_vars["max_gen"].get()),
            crossover_prob=float(self._param_vars["crossover"].get()),
            mutation_rate_base=float(self._param_vars["mutation"].get()),
            tournament_size=int(self._param_vars["tournament"].get()),
            local_search_iters=int(self._param_vars["ls_iters"].get()),
            tabu_tenure=int(self._param_vars["tabu_tenure"].get()),
            stagnation_limit=int(self._param_vars["stagnation"].get()),
            dual_decode=self.use_dual_decode_var.get(),
            use_psgs_in_ls=self.use_psgs_var.get(),
            time_limit_sec=float(self._param_vars["time_limit"].get()),
            verbose=False,
            seed=None,  # random seed for extra diversity
        )

        for idx, fpath in enumerate(files):
            if self._stop_event.is_set():
                self._log("⏹ Solver stopped by user.", "warning")
                break

            name = os.path.splitext(os.path.basename(fpath))[0]
            self._log(f"\n── Instance {idx+1}/{total}: {name} ──", "accent")

            try:
                inst = parse_instance(fpath)
                self._log(f"   Tâches: {inst.n_jobs}  Ressources: {inst.n_resources}  "
                          f"Horizon: {inst.horizon}", "info")

                bounds = compute_lower_bound(inst)
                lb = bounds['LB']
                self._log(f"   LB(CPM)={bounds['LB1']}  LB(Res)={bounds['LB2']}  "
                          f"LB(Window)={bounds['LB3']}  → LB={lb}", "info")

                # Solutions de référence
                ref_lb = self.solutions_ref.get(name.lower(), (None, None))[0]
                ref_ub = self.solutions_ref.get(name.lower(), (None, None))[1]
                if ref_ub:
                    self._log(f"   Reference: LB={ref_lb}  UB={ref_ub}", "info")

                def progress_cb(gen, max_gen, best_cmax, elapsed):
                    pct = gen / max(1, max_gen) * 100
                    self.progress_var.set(pct)
                    self.progress_label_var.set(
                        f"Gen {gen}/{max_gen} — Best Cmax: {best_cmax}"
                        + (f"  (Ref UB: {ref_ub})" if ref_ub else "")
                        + f"  [{elapsed:.0f}s]")
                    if gen % 50 == 0:
                        gap = (best_cmax - lb) / lb * 100 if lb > 0 else 0
                        self._log(f"   Gen {gen}: Cmax={best_cmax}  Gap={gap:.1f}%  [{elapsed:.0f}s]", "info")

                t0 = time.time()
                if name.startswith("j60"):
                    self._log(f"   [Fast Mode] Using ultra-fast j60 solver (Budget: {config.time_limit_sec}s)...", "info")
                    cmax, schedule = fast_solve(inst, time_limit=config.time_limit_sec, lb=lb, ub_target=ref_ub)
                else:
                    solver = HybridGASolver(inst, config, progress_cb=progress_cb,
                                            stop_event=self._stop_event)
                    schedule = solver.solve()
                elapsed = time.time() - t0

                cmax = schedule.cmax
                gap_lb = ((cmax - lb) / lb * 100) if lb > 0 else 0
                gap_ub = ((cmax - ref_ub) / ref_ub * 100) if ref_ub else 0
                beats = ref_ub is not None and cmax < ref_ub

                result = SolverResult(
                    instance_name=name,
                    n_jobs=inst.n_jobs,
                    our_cmax=cmax,
                    lb=lb,
                    ref_lb=ref_lb,
                    ref_ub=ref_ub,
                    gap_to_lb=gap_lb,
                    gap_to_ub=gap_ub,
                    elapsed=elapsed,
                    schedule=schedule,
                    beats_known=beats,
                )
                self.results.append(result)
                self.current_schedule = schedule
                self._instances_cache = getattr(self, '_instances_cache', {})
                self._instances_cache[name] = inst

                status = "🏆 BEATS KNOWN!" if beats else ("✓ Open" if not ref_ub or ref_lb != ref_ub else "✓")
                log_level = "success" if beats else "success"
                self._log(
                    f"   ✓ Cmax={cmax}  Gap-LB={gap_lb:.1f}%"
                    + (f"  Gap-UB={gap_ub:.1f}%" if ref_ub else "")
                    + f"  Time={elapsed:.1f}s  {status}",
                    log_level)

                # Update UI in main thread
                self.after(0, self._add_result_row, result)
                self.after(0, self._refresh_gantt)

            except Exception as e:
                self._log(f"   ✗ Erreur on {name}: {e}", "error")

            pct = (idx + 1) / total * 100
            self.progress_var.set(pct)

        self._log("\n✅ Solveur terminé.", "success")
        self.progress_label_var.set("Terminé")
        self.after(0, self._update_summary_bar)

    # ─── Gantt Chart ──────────────────────────────────────────────────────

    def _refresh_gantt(self):
        if not HAS_MATPLOTLIB:
            return
        sched = self.current_schedule
        if sched is None:
            return
        inst = sched.instance
        self._draw_gantt(sched, inst)

    def _draw_gantt(self, sched: Schedule, inst: RCPSPInstance):
        fig = self.gantt_fig
        ax = self.gantt_ax
        ax.clear()
        fig.set_facecolor(BG_CARD)
        ax.set_facecolor(BG_PANEL)

        n = inst.n_jobs
        ft = sched.finish_times
        cmax = sched.cmax

        real_jobs = list(range(n))
        # Draw bars
        for i in real_jobs:
            s = sched.start_times[i]
            d = inst.durations[i]
            if d == 0:
                continue
            color = GANTT_COLORS[i % len(GANTT_COLORS)]
            bar = ax.barh(y=i, width=d, left=s, height=0.6,
                          color=color, alpha=0.88,
                          edgecolor=BG_CARD, linewidth=0.5)
            # Label inside bar if wide enough
            if d >= 2:
                ax.text(s + d / 2, i, f"J{i}", ha='center', va='center',
                        fontsize=7, color='white', fontweight='bold')

        # Vertical line at Cmax
        ax.axvline(x=cmax, color=DANGER, linewidth=1.5, linestyle='--', alpha=0.8)
        ax.text(cmax + 0.2, n * 0.98, f'Cmax={cmax}',
                color=DANGER, fontsize=9, fontweight='bold', va='top')

        ax.set_xlabel("Time", color=TEXT_DIM)
        ax.set_ylabel("Job Index", color=TEXT_DIM)
        ax.set_title(f"Gantt Chart — {inst.name}  (Cmax={cmax})",
                     color=TEXT_MAIN, fontsize=12, fontweight='bold', pad=10)
        ax.tick_params(colors=TEXT_DIM)
        ax.spines['bottom'].set_color(BORDER)
        ax.spines['left'].set_color(BORDER)
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
        ax.set_xlim(0, max(cmax + 5, 10))
        ax.set_ylim(-0.5, n - 0.5)
        ax.grid(axis='x', color=BORDER, alpha=0.5, linewidth=0.5)
        fig.tight_layout()
        self.gantt_canvas.draw()

        feasible = sched.is_feasible()
        self.gantt_info_var.set(
            f"Instance: {inst.name}  |  Tâches: {n}  |  Cmax: {cmax}  "
            f"|  Feasible: {'✅' if feasible else '❌'}")

    def _save_gantt_png(self):
        if self.current_schedule is None:
            messagebox.showwarning("No chart", "No schedule to save.")
            return
        path = filedialog.asksaveasfilename(
            defaultextension=".png",
            filetypes=[("PNG Image", "*.png"), ("All files", "*.*")],
            initialfile=f"gantt_{self.current_schedule.instance.name}.png")
        if path:
            self.gantt_fig.savefig(path, dpi=150, bbox_inches='tight',
                                   facecolor=BG_CARD)
            self._log(f"Gantt chart saved to {path}", "success")

    # ─── Results Table ────────────────────────────────────────────────────

    def _add_result_row(self, r: SolverResult):
        gap_ub_str = f"{r.gap_to_ub:.1f}%" if r.ref_ub else "—"
        ref_lb_str = str(r.ref_lb) if r.ref_lb else "—"
        ref_ub_str = str(r.ref_ub) if r.ref_ub else "—"
        open_inst  = r.ref_ub and r.ref_lb and r.ref_lb < r.ref_ub

        status = "🏆 BEATS KNOWN!" if r.beats_known else ("⚡ Open" if open_inst else "✓ Solved")
        tag = "beat" if r.beats_known else ("open" if open_inst else "neutral")

        self.results_tree.insert("", "end", values=(
            r.instance_name,
            r.n_jobs,
            r.our_cmax,
            r.lb,
            ref_lb_str,
            ref_ub_str,
            f"{r.gap_to_lb:.1f}%",
            gap_ub_str,
            f"{r.elapsed:.1f}",
            status,
        ), tags=(tag,))
        self._update_summary_bar()

    def _clear_results(self):
        self.results.clear()
        for item in self.results_tree.get_children():
            self.results_tree.delete(item)
        self._update_summary_bar()

    def _update_summary_bar(self):
        for widget in self.summary_frame.winfo_children():
            widget.destroy()

        total = len(self.results)
        beats = sum(1 for r in self.results if r.beats_known)
        open_inst = sum(1 for r in self.results if r.ref_ub and r.ref_lb and r.ref_lb < r.ref_ub)
        avg_gap = (sum(r.gap_to_lb for r in self.results) / total) if total > 0 else 0

        def _stat(frame, value, label, color):
            f = tk.Frame(frame, bg=BG_CARD, padx=16, pady=6)
            f.pack(side="left", padx=4)
            tk.Label(f, text=str(value), font=("Segoe UI", 18, "bold"),
                     bg=BG_CARD, fg=color).pack()
            tk.Label(f, text=label, font=FONT_SMALL,
                     bg=BG_CARD, fg=TEXT_DIM).pack()

        _stat(self.summary_frame, total, "Solved", ACCENT)
        _stat(self.summary_frame, beats, "Beat Known UB", ACCENT2)
        _stat(self.summary_frame, open_inst, "Open Instances", WARNING)
        _stat(self.summary_frame, f"{avg_gap:.1f}%", "Avg Gap to LB", TEXT_MAIN)

    # ─── Export ───────────────────────────────────────────────────────────

    def _export_excel(self):
        if not self.results:
            messagebox.showwarning("No results", "No results to export.")
            return
        try:
            import openpyxl
            from openpyxl.styles import Font, PatternFill, Alignment
        except ImportErreur:
            messagebox.showerror("Missing library",
                                 "openpyxl not installed.\nRun: pip install openpyxl")
            return

        path = filedialog.asksaveasfilename(
            defaultextension=".xlsx",
            filetypes=[("Excel", "*.xlsx")],
            initialfile="rcpsp_results.xlsx")
        if not path:
            return

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "RCPSP Results"

        headers = ["Instance", "Jobs", "Our Cmax", "LB", "Ref LB", "Ref UB",
                   "Gap to LB %", "Gap to UB %", "Time (s)", "Status"]
        header_font = Font(bold=True, color="FFFFFF")
        header_fill = PatternFill("solid", fgColor="6C63FF")

        for col, h in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col, value=h)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal="center")

        for row_idx, r in enumerate(self.results, 2):
            open_inst = r.ref_ub and r.ref_lb and r.ref_lb < r.ref_ub
            status = "BEATS KNOWN!" if r.beats_known else ("Open" if open_inst else "Solved")
            ws.append([
                r.instance_name, r.n_jobs, r.our_cmax, r.lb,
                r.ref_lb or "", r.ref_ub or "",
                round(r.gap_to_lb, 2), round(r.gap_to_ub, 2) if r.ref_ub else "",
                round(r.elapsed, 2), status,
            ])
            if r.beats_known:
                fill = PatternFill("solid", fgColor="1A3320")
                for col in range(1, len(headers) + 1):
                    ws.cell(row=row_idx, column=col).fill = fill

        for col in ws.columns:
            ws.column_dimensions[col[0].column_letter].width = 16

        wb.save(path)
        self._log(f"✓ Results exported to {path}", "success")
        messagebox.showinfo("Export", f"Results saved to:\n{path}")

    def _export_word(self):
        if not self.results:
            messagebox.showwarning("No results", "No results to export.")
            return
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportErreur:
            messagebox.showerror("Missing library",
                                 "python-docx not installed.\nRun: pip install python-docx")
            return

        path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Document", "*.docx")],
            initialfile="rcpsp_results.docx")
        if not path:
            return

        doc = Document()
        doc.add_heading("RCPSP Results — SupNum S3C'1447", level=0)
        doc.add_heading("Hybrid Genetic Algorithm + Neighborhood Search", level=1)
        doc.add_paragraph(
            f"Total instances solved: {len(self.results)}\n"
            f"Instances beating known UB: {sum(1 for r in self.results if r.beats_known)}")

        table = doc.add_table(rows=1, cols=9)
        table.style = "Table Grid"
        hdr_row = table.rows[0].cells
        headers = ["Instance", "Jobs", "Our Cmax", "LB", "Ref LB", "Ref UB",
                   "Gap LB%", "Gap UB%", "Status"]
        for i, h in enumerate(headers):
            hdr_row[i].text = h
            hdr_row[i].paragraphs[0].runs[0].bold = True

        for r in self.results:
            open_inst = r.ref_ub and r.ref_lb and r.ref_lb < r.ref_ub
            status = "BEATS KNOWN!" if r.beats_known else ("Open" if open_inst else "Solved")
            row = table.add_row().cells
            vals = [
                r.instance_name, str(r.n_jobs), str(r.our_cmax), str(r.lb),
                str(r.ref_lb or "—"), str(r.ref_ub or "—"),
                f"{r.gap_to_lb:.1f}%",
                f"{r.gap_to_ub:.1f}%" if r.ref_ub else "—",
                status,
            ]
            for i, v in enumerate(vals):
                row[i].text = v

        doc.save(path)
        self._log(f"✓ Word report saved to {path}", "success")
        messagebox.showinfo("Export", f"Word report saved to:\n{path}")

    # ─── Helpers ──────────────────────────────────────────────────────────

    def _log(self, msg: str, level: str = "info"):
        def _do():
            ts = time.strftime("%H:%M:%S")
            self.log_text.configure(state="normal")
            self.log_text.insert(tk.END, f"[{ts}] {msg}\n", level)
            self.log_text.see(tk.END)
            self.log_text.configure(state="disabled")
        self.after(0, _do)

    def _clear_log(self):
        self.log_text.configure(state="normal")
        self.log_text.delete("1.0", tk.END)
        self.log_text.configure(state="disabled")

    def _make_card(self, parent, title: str) -> tk.Frame:
        outer = tk.Frame(parent, bg=BORDER, padx=1, pady=1)
        inner = tk.Frame(outer, bg=BG_CARD, padx=12, pady=10)
        inner.pack(fill="both", expand=True)
        tk.Label(inner, text=title, font=("Segoe UI", 10, "bold"),
                 bg=BG_CARD, fg=ACCENT).pack(anchor="w", pady=(0, 6))
        outer.pack(fill="x", expand=False)
        return inner

    def _btn(self, parent, text, command, color=None):
        c = color or ACCENT
        b = tk.Button(parent, text=text, command=command,
                      bg=c, fg="white", relief="flat",
                      font=("Segoe UI", 9, "bold"),
                      padx=10, pady=5,
                      activebackground=c, activeforeground="white",
                      cursor="hand2")
        b.bind("<Enter>", lambda e: b.config(bg=self._lighten(c)))
        b.bind("<Leave>", lambda e: b.config(bg=c))
        return b

    @staticmethod
    def _lighten(hex_color: str) -> str:
        """Lighten a hex color by 20%."""
        try:
            hex_color = hex_color.lstrip('#')
            r, g, b = int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16)
            r = min(255, int(r * 1.2))
            g = min(255, int(g * 1.2))
            b = min(255, int(b * 1.2))
            return f'#{r:02x}{g:02x}{b:02x}'
        except Exception:
            return hex_color

    def _style_axes(self, ax, title=""):
        ax.set_facecolor(BG_PANEL)
        ax.set_title(title, color=TEXT_DIM, fontsize=11)
        ax.tick_params(colors=TEXT_DIM)
        ax.spines['bottom'].set_color(BORDER)
        ax.spines['left'].set_color(BORDER)
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)

    def _on_close(self):
        self._stop_event.set()
        self.destroy()
